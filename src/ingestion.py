"""多格式数据摄入模块：Excel / PDF / Word / 图片 → 统一 DataFrame

标准输出列（与 processor.py 约定的考研初试数据格式对齐）：
    姓名, 总分, 政治, 英语, 业务课一, 业务课二
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pandas as pd

# 统一标准列名
STANDARD_COLUMNS = ["姓名", "总分", "政治", "英语", "业务课一", "业务课二"]

# ── 语义列名映射表（标准键 → 常见别名列表）─────────────────────────
COLUMN_MAP: dict[str, list[str]] = {
    "政治": ["政治", "政治理论", "政治类", "思想政治", "思想政治理论",
            "politics", "政治成绩", "政治分"],
    "英语": ["英语", "外国语", "外语", "英语一", "英语二", "english",
            "外语成绩", "英语成绩"],
    "业务课一": ["业务课一", "业务课1", "业务课1成绩", "专一", "专业课一",
               "专业一", "专业1", "科目一", "初试专业课一",
               "subject1", "数学", "业务1", "业务一"],
    "业务课二": ["业务课二", "业务课2", "业务课2成绩", "专二", "专业课二",
               "专业二", "专业2", "科目二", "初试专业课二",
               "subject2", "业务2", "业务二"],
    "姓名": ["姓名", "name", "考生姓名", "考生", "学生姓名", "candidate",
            "candidatename", "名字", "考生名称"],
    "总分": ["总分", "total", "总成绩", "总得分", "total_score",
            "totalscore", "sum", "初试总分", "初试总成绩", "考试总分"],
}
# 反向索引：所有别名的扁平集（用于快速精确匹配）
_ALIAS_FLAT: dict[str, str] = {}
for std, aliases in COLUMN_MAP.items():
    for a in aliases:
        _ALIAS_FLAT[a] = std
        _ALIAS_FLAT[a.lower().replace(" ", "")] = std

# 启发式列名映射表（输入列名 → 标准列名）
COLUMN_ALIASES: dict[str, str] = {
    # 姓名
    "姓名": "姓名", "name": "姓名", "考生姓名": "姓名", "考生": "姓名",
    "学生姓名": "姓名", "candidate": "姓名", "candidatename": "姓名",
    # 总分
    "总分": "总分", "total": "总分", "总成绩": "总分", "总得分": "总分",
    "total_score": "总分", "totalscore": "总分", "sum": "总分",
    "初试总分": "总分", "初试总成绩": "总分", "考试总分": "总分",
    # 政治
    "政治": "政治", "政治理论": "政治", "politics": "政治",
    "思想政治": "政治", "思想政治理论": "政治",
    # 英语
    "英语": "英语", "english": "英语", "外国语": "英语",
    "外语": "英语", "英语一": "英语", "英语二": "英语",
    # 业务课一
    "业务课一": "业务课一", "专一": "业务课一", "专业课一": "业务课一",
    "subject1": "业务课一", "专业一": "业务课一", "数学": "业务课一",
    "业务1": "业务课一", "科目一": "业务课一",
    "业务课1成绩": "业务课一", "业务课1": "业务课一",
    # 业务课二
    "业务课二": "业务课二", "专二": "业务课二", "专业课二": "业务课二",
    "subject2": "业务课二", "专业二": "业务课二", "业务2": "业务课二",
    "科目二": "业务课二",
    "业务课2成绩": "业务课二", "业务课2": "业务课二",
}


# ── 学校名校对映射表 ────────────────────────────────────────────────
SCHOOL_CN_MAP = {
    "Nanjing_Normal_University": "南京师范大学",
    "Fudan_University": "复旦大学",
    "Peking_University": "北京大学",
    "Tsinghua_University": "清华大学",
    "Zhejiang_University": "浙江大学",
    "Sun_Yat_sen_University": "中山大学",
    "Wuhan_University": "武汉大学",
    "Sichuan_University": "四川大学",
    "Xiamen_University": "厦门大学",
    "China_Agricultural_University": "中国农业大学",
    "Huazhong_University_of_Science_and_Technology": "华中科技大学",
    "Harbin_Institute_of_Technology": "哈尔滨工业大学",
    "Shandong_University": "山东大学",
    "Xi_an_Jiaotong_University": "西安交通大学",
    "Shanghai_Jiao_Tong_University": "上海交通大学",
}

_SCHOOL_CN_KEYWORDS: dict[str, str] = {}
for _en, _cn in SCHOOL_CN_MAP.items():
    _SCHOOL_CN_KEYWORDS[_en] = _cn
    for _part in _en.replace("_", " ").split():
        _SCHOOL_CN_KEYWORDS[_part.lower()] = _cn


def extract_school_name(filename: str) -> str:
    """从文件名中自动提取学校中文名.

    优先级：
        1. 在 ``SCHOOL_CN_MAP`` 的英文键中精确匹配前缀
        2. 在 ``SCHOOL_CN_MAP`` 的中文值中匹配
        3. 取文件名中第一个独立单词（大写开头）作为英文名兜底
    """
    stem = Path(filename).stem  # 去后缀

    # 1. 英文键前缀匹配（去掉下划线后逐个检查）
    for en_key, cn_name in SCHOOL_CN_MAP.items():
        if stem.startswith(en_key):
            return cn_name

    # 2. 中文校名直接出现在文件名中
    for cn_name in SCHOOL_CN_MAP.values():
        if cn_name in stem:
            return cn_name

    # 3. 取第一个大写单词
    import re as _re
    words = _re.findall(r"[A-Z][a-z]+", stem)
    if words:
        guess = words[0]
        for ek, cn in SCHOOL_CN_MAP.items():
            if guess.lower() in ek.lower().replace("_", ""):
                return cn

    return "目标高校"


def _ocr_ppstructure(image_input: str | Path | Any) -> pd.DataFrame | None:
    """尝试用 PaddleOCR PPStructure (v3) 识别图片中的表格.

    Returns
    -------
    pd.DataFrame | None
        成功返回 DataFrame，失败返回 None.
    """
    try:
        from paddleocr import PPStructureV3
        engine = PPStructureV3(lang="ch")
        result = engine.predict(str(image_input))
        # PPStructureV3 返回生成器 / 列表
        frames: list[pd.DataFrame] = []
        for item in result if isinstance(result, (list, tuple)) else [result]:
            html = getattr(item, "html", None) or (isinstance(item, dict) and item.get("html"))
            if html:
                tables = pd.read_html(html)
                for tbl in tables:
                    if not tbl.empty:
                        tbl.columns = [str(c).strip().replace("\n", "") for c in tbl.columns]
                        frames.append(tbl)
        return pd.concat(frames, ignore_index=True) if frames else None
    except ImportError:
        return None
    except Exception:
        return None


def _ocr_tesseract(image_input: str | Path | Any) -> pd.DataFrame | None:
    """用 Tesseract OCR 提取图片文字，尝试构建 DataFrame.

    由于 Tesseract 无表格结构识别能力，仅能提取纯文本。
    返回包含单列 ``text`` 的 DataFrame，或 None.
    """
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(image_input) if isinstance(image_input, (str, Path)) else image_input
        data = pytesseract.image_to_string(img, lang="chi_sim+eng")
        lines = [line.strip() for line in data.split("\n") if line.strip()]
        if not lines:
            return None
        # 尝试按分隔符拆分
        rows: list[list[str]] = []
        for line in lines:
            parts = [p.strip() for p in line.split() if p.strip()]
            if parts:
                rows.append(parts)
        return pd.DataFrame(rows) if rows else None
    except ImportError:
        return None
    except Exception:
        return None


def preprocess_image(image_path: str | Path) -> pd.DataFrame:
    """OCR 识别图片中的表格并转换为 Pandas DataFrame.

    优先级：
        1. PPStructureV3（需 paddleocr + paddlepaddle）
        2. pytesseract（轻量，已安装）

    Returns
    -------
    pd.DataFrame
        识别结果.
    """
    path = Path(image_path)
    if not path.exists():
        raise FileNotFoundError(f"图片文件不存在: {path}")

    # 1. PPStructureV3
    df = _ocr_ppstructure(path)
    if df is not None and not df.empty:
        return df

    # 2. pytesseract 兜底
    df = _ocr_tesseract(path)
    if df is not None and not df.empty:
        return df

    raise ValueError(
        f"OCR 未能从图片中提取出有效表格: {image_path}\n"
        f"  已尝试: PPStructureV3 / Tesseract\n"
        f"  提示: 如需 PaddleOCR 表格识别，请安装 pip install paddlepaddle paddleocr"
    )


def normalize_columns(
    df: pd.DataFrame,
    verbose: bool = True,
) -> pd.DataFrame:
    """语义列名模糊映射引擎.

    将 DataFrame 的列名与 ``COLUMN_MAP`` 中的别名做精确 + 模糊匹配，
    映射为标准列（政治/英语/业务课一/业务课二/姓名/总分），
    其余列重命名为 ``raw_{原名}``。

    如果四门核心科目（政治/英语/业务课一/业务课二）有缺失，
    在 DataFrame 的 attrs 中设置 ``_missing_core`` 标记。

    Parameters
    ----------
    df : pd.DataFrame
    verbose : bool
        是否打印匹配日志.

    Returns
    -------
    pd.DataFrame
    """
    print(f"DEBUG: 传入的原始列名: {df.columns.tolist()}")

    if df.empty:
        return df

    df = df.copy()
    # 表头预处理：去换行/空格/括号内容，独立数字1/2 → 一/二
    def _clean_col(c: object) -> str:
        s = str(c).strip().replace("\n", "").replace(" ", "").replace("\u3000", "")
        s = re.sub(r"[（(][^）)]*[）)]", "", s)  # 删除 (100分) 等后缀
        s = re.sub(r"(?<!\d)(\d)(?!\d)", lambda m: {"1": "一", "2": "二"}.get(m.group(1), m.group(1)), s)
        return s

    raw_names = {c: _clean_col(c) for c in df.columns}
    df.columns = list(raw_names.values())
    reverse_map: dict[str, str] = {}  # cleaned_name → standard_key
    print(f"  [DEBUG] 清洗后列名: {list(df.columns)}")

    # ── 1. 精确匹配（直接查 _ALIAS_FLAT） ────────────────────────────
    for col in df.columns:
        if col in _ALIAS_FLAT:
            reverse_map[col] = _ALIAS_FLAT[col]
            continue
        lowered = col.lower()
        if lowered in _ALIAS_FLAT:
            reverse_map[col] = _ALIAS_FLAT[lowered]
            continue

    # ── 2. 模糊匹配（difflib） ────────────────────────────────────────
    import difflib

    ALL_ALIASES = sorted(set(_ALIAS_FLAT.keys()), key=len, reverse=True)
    for col in df.columns:
        if col in reverse_map:
            continue
        matches = difflib.get_close_matches(col, ALL_ALIASES, n=3, cutoff=0.80)
        if not matches:
            matches = difflib.get_close_matches(col.lower(), ALL_ALIASES, n=3, cutoff=0.80)
        if matches:
            best = matches[0]
            target = _ALIAS_FLAT.get(best) or _ALIAS_FLAT.get(best.lower())
            # 禁止将非姓名列强行映射为"姓名"
            if target == "姓名" and "姓名" not in col and "name" not in col.lower():
                if verbose:
                    print(f"  [FUZZY] 列「{col}」相似 {best} 但拒绝映射为「姓名」（置信度不足）")
                continue
            reverse_map[col] = target
            if verbose and len(matches) > 1:
                candidates = [f"{m}→{_ALIAS_FLAT.get(m, m)}" for m in matches[:3]]
                print(f"  [FUZZY] 列「{col}」模糊匹配: {candidates}")

    # ── 3. 重命名 ─────────────────────────────────────────────────────
    rename: dict[str, str] = {}
    for col in df.columns:
        if col in reverse_map:
            rename[col] = reverse_map[col]
        else:
            rename[col] = f"raw_{col}"

    df = df.rename(columns=rename)

    # ── 4. 检查核心科目是否齐全 ──────────────────────────────────────
    core = {"政治", "英语", "业务课一", "业务课二"}
    missing = core - set(df.columns)
    if missing:
        msg = f"未发现核心科目列: {missing}，请在侧边栏【数据结构修正】面板中手动指定列名。"
        if verbose:
            print(f"  [WARN] {msg}")
        df.attrs["_missing_core"] = msg
    else:
        df.attrs["_missing_core"] = ""

    return df


class DataIngestor:
    """智能数据摄入器，支持多种格式 → 统一 DataFrame.

    无论输入是 Excel / PDF / Word / 图片，最终输出列均为：
        ``["姓名", "总分", "政治", "英语", "业务课一", "业务课二"]``
    无法映射的原始列保留在原位，不会被丢弃。
    """

    # ── 公共入口 ──────────────────────────────────────────────────────

    def detect_and_load(self, file_path: str | Path) -> pd.DataFrame:
        """根据文件后缀自动选择对应的加载方法.

        Parameters
        ----------
        file_path : str | Path

        Returns
        -------
        pd.DataFrame
            标准化后的 DataFrame.

        Raises
        ------
        FileNotFoundError
            文件不存在.
        ValueError
            不支持的文件格式.
        """
        path = Path(file_path)
        print(f"DEBUG: 正在加载文件: {file_path}")
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()

        loader_map: dict[str, Any] = {
            ".xlsx": self.load_excel,
            ".xls":  self.load_excel,
            ".pdf":  self.load_pdf,
            ".docx": self.load_word,
            ".doc":  self.load_word,
            ".jpg":  self.load_image,
            ".jpeg": self.load_image,
            ".png":  self.load_image,
            ".bmp":  self.load_image,
            ".tiff": self.load_image,
        }

        loader = loader_map.get(suffix)
        if loader is None:
            raise ValueError(
                f"不支持的文件格式: {suffix}\n"
                f"  支持的类型: xlsx / xls / pdf / docx / doc / jpg / jpeg / png / bmp / tiff"
            )

        try:
            return loader(path)
        except Exception as e:
            raise RuntimeError(f"加载文件失败: {path}\n  错误信息: {e}") from e

    def load_data_dir(self, data_dir: str | Path = "data") -> pd.DataFrame:
        """批量加载 data/ 目录下所有支持的文件，纵向合并.

        Parameters
        ----------
        data_dir : str | Path
            数据目录路径，默认为 ``data/``.

        Returns
        -------
        pd.DataFrame

        Raises
        ------
        FileNotFoundError
            目录不存在.
        ValueError
            目录中无可识别文件.
        """
        dir_path = Path(data_dir)
        if not dir_path.exists():
            raise FileNotFoundError(f"数据目录不存在: {dir_path}")

        supported = {
            ".xlsx", ".xls", ".pdf", ".docx", ".doc",
            ".jpg", ".jpeg", ".png", ".bmp", ".tiff",
        }
        frames: list[pd.DataFrame] = []

        for f in sorted(dir_path.iterdir()):
            if f.suffix.lower() not in supported:
                continue
            try:
                df = self.detect_and_load(f)
                frames.append(df)
            except Exception as exc:
                print(f"[WARN] 跳过 {f.name} —— {exc}")

        if not frames:
            raise ValueError(
                f"目录 {data_dir} 下未找到可识别的数据文件\n"
                f"  支持的格式: {', '.join(sorted(supported))}"
            )

        return pd.concat(frames, ignore_index=True)

    # ── Excel 适配器 ──────────────────────────────────────────────────

    @staticmethod
    def _detect_header_row(
        file_path: str | Path,
        header_keywords: tuple[str, ...] | None = None,
    ) -> int:
        """扫描 Excel 前 rows_to_scan 行，返回表头所在行号（0-indexed）.

        找不到时返回 0（默认第一行）.
        """
        if header_keywords is None:
            header_keywords = ("姓名", "考生编号", "考生", "总分", "总成绩",
                               "专业", "政治", "英语", "外语", "业务课")
        rows_to_scan = 5
        try:
            preview = pd.read_excel(file_path, nrows=rows_to_scan, header=None)
        except Exception:
            return 0
        for i in range(min(rows_to_scan, len(preview))):
            row_vals = [str(c).strip() for c in preview.iloc[i].fillna("")]
            row_text = " ".join(row_vals)
            hits = sum(1 for kw in header_keywords if kw in row_text)
            if hits >= 2:
                print(f"  [HEADER] 第 {i+1} 行被检测为表头（命中 {hits} 个关键词）")
                return i
        return 0

    def load_excel(self, file_path: str | Path) -> pd.DataFrame:
        """读取 Excel（xlsx / xls）并标准化.

        自动扫描前 5 行定位表头，非第一行表头也能正确识别。

        Raises
        ------
        ValueError
            Excel 解析失败或内容为空.
        """
        try:
            header_row = self._detect_header_row(file_path)
            df = pd.read_excel(file_path, header=header_row)
            print(f"  [DEBUG] 使用第 {header_row + 1} 行作为表头")
        except Exception as e:
            raise ValueError(
                f"Excel 读取失败: {file_path}\n"
                f"  请确认文件是有效的 Excel 格式且未被损坏。\n"
                f"  错误详情: {e}"
            ) from e

        if df.empty:
            raise ValueError(f"Excel 文件内容为空: {file_path}")

        # 调试输出
        print("  [DEBUG] df.head(5):")
        print(df.head(5).to_string())
        print(f"  [DEBUG] df.columns: {list(df.columns)}")

        # 丢弃 Unnamed NaN 列
        df = self._drop_unnamed_columns(df)

        return self._normalize(df)

    @staticmethod
    def _drop_unnamed_columns(df: pd.DataFrame) -> pd.DataFrame:
        """删除列名包含 Unnamed 且数据全为 NaN 的列."""
        to_drop = []
        for col in df.columns:
            if "Unnamed" in str(col) or "unnamed" in str(col):
                if df[col].isna().all():
                    to_drop.append(col)
        if to_drop:
            print(f"  [CLEAN] 删除 {len(to_drop)} 个 Unnamed 空列: {to_drop}")
            df = df.drop(columns=to_drop)
        return df

    # ── PDF 适配器 ────────────────────────────────────────────────────

    def load_pdf(self, file_path: str | Path) -> pd.DataFrame:
        """使用 pdfplumber 提取 PDF 中的表格.

        通过 ``table_settings`` 参数处理复杂表格的对齐问题：
          - vertical_strategy="text"   : 根据文字边缘检测列边界
          - horizontal_strategy="text" : 根据文字边缘检测行边界
          - keep_blank_chars=True      : 保留空白单元格占位

        Raises
        ------
        ValueError
            PDF 中未提取到有效表格.
        """
        try:
            import pdfplumber
        except ImportError as e:
            raise ImportError("pdfplumber 未安装，请执行: pip install pdfplumber") from e

        frames: list[pd.DataFrame] = []

        table_settings = {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "intersection_x_tolerance": 5,
            "intersection_y_tolerance": 5,
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                saved_header: list[str] | None = None
                for page in pdf.pages:
                    tables = page.extract_tables(table_settings)
                    if tables:
                        for table in tables:
                            if not table or len(table) < 2:
                                continue
                            first_row = [self._clean_cell(c) for c in table[0]]
                            if _looks_like_header(first_row):
                                saved_header = first_row
                                rows_data = table[1:]
                            else:
                                rows_data = table
                            rows = []
                            for row in rows_data:
                                rows.append([self._clean_cell(c) for c in row])
                            if saved_header is None:
                                saved_header = [f"col_{i}" for i in range(len(rows[0]))]
                            pdf_df = pd.DataFrame(rows, columns=saved_header[:len(rows[0])])
                            frames.append(pdf_df)
                    else:
                        # 无表格 → 尝试自由文本提取
                        text = page.extract_text()
                        if text:
                            text_df = _parse_free_text(text)
                            if text_df is not None:
                                frames.append(text_df)
        except Exception as e:
            raise ValueError(
                f"PDF 解析失败: {file_path}\n"
                f"  请确认文件是有效的 PDF 且包含可提取的表格。\n"
                f"  错误详情: {e}"
            ) from e

        # ── pdfplumber 失败 → PaddleOCR 兜底 ──────────────────────────
        if not frames:
            try:
                ocr_frames = self._ocr_pdf(file_path)
                frames.extend(ocr_frames)
            except ImportError:
                raise ImportError(
                    "PDF 无法提取表格且 PaddleOCR 未安装，请执行: pip install paddleocr"
                )
            except Exception as e:
                raise ValueError(
                    f"PDF 解析失败（含 OCR 兜底）: {file_path}\n  错误: {e}"
                ) from e

        if not frames:
            raise ValueError(
                f"PDF 中未提取到有效表格: {file_path}\n"
                f"  请确认 PDF 中包含可解析的表格结构，或安装 PaddleOCR 识别扫描件。"
            )

        result = pd.concat(frames, ignore_index=True)
        result = self._drop_unnamed_columns(result)
        print(f"  [DEBUG] PDF 提取结果: {result.shape}")
        print(f"  [DEBUG] df.columns: {list(result.columns)}")
        return self._normalize(result)

    # ── PDF 兜底：PaddleOCR ────────────────────────────────────────────

    def _ocr_pdf(self, file_path: str | Path) -> list[pd.DataFrame]:
        """将 PDF 各页转为图片后使用 PaddleOCR PPStructureV3 提取表格.

        适用于扫描件 / 图片型 PDF，pdfplumber 无法提取时自动调用。

        Returns
        -------
        list[pd.DataFrame]
            每页识别出的表格 DataFrame 列表.
        """
        from pdf2image import convert_from_path

        frames: list[pd.DataFrame] = []
        images = convert_from_path(file_path, dpi=300)

        # 尝试 PPStructureV3（需要 paddlepaddle）
        try:
            from paddleocr import PPStructureV3
            engine = PPStructureV3(lang="ch")
            for page_idx, image in enumerate(images):
                import tempfile, os
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp_path = tmp.name
                    image.save(tmp_path, "PNG")
                try:
                    result = engine.predict(tmp_path)
                    for item in result if isinstance(result, (list, tuple)) else [result]:
                        html = getattr(item, "html", None) or (
                            isinstance(item, dict) and item.get("html"))
                        if html:
                            tables = pd.read_html(html)
                            for tbl in tables:
                                if not tbl.empty:
                                    tbl.columns = [str(c).strip().replace("\n", "") for c in tbl.columns]
                                    frames.append(tbl)
                except Exception as e:
                    print(f"  [WARN] PPStructureV3 第 {page_idx+1} 页失败: {e}")
                finally:
                    os.unlink(tmp_path)
            if frames:
                return frames
        except ImportError:
            pass

        # 兜底：使用 _ocr_ppstructure（尝试各种 PaddleOCR 兼容接口）
        for page_idx, image in enumerate(images):
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
                image.save(tmp_path, "PNG")
            try:
                df = _ocr_ppstructure(tmp_path)
                if df is not None and not df.empty:
                    frames.append(df)
            except Exception:
                pass
            finally:
                os.unlink(tmp_path)

        return frames

    # ── Word 适配器 ───────────────────────────────────────────────────

    def load_word(self, file_path: str | Path) -> pd.DataFrame:
        """使用 python-docx 提取 Word 文档中的表格.

        Raises
        ------
        ValueError
            Word 中未提取到有效表格.
        """
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError("python-docx 未安装，请执行: pip install python-docx") from e

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(
                f"Word 文档读取失败: {file_path}\n"
                f"  请确认文件是有效的 .docx 格式。\n"
                f"  错误详情: {e}"
            ) from e

        frames: list[pd.DataFrame] = []

        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)
            if len(rows_data) >= 2:
                header = rows_data[0]
                data = rows_data[1:]
                doc_df = pd.DataFrame(data, columns=header)
                frames.append(doc_df)

        if not frames:
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                text_df = _parse_free_text("\n".join(paragraphs))
                if text_df is not None:
                    frames.append(text_df)

        if not frames:
            raise ValueError(
                f"Word 文档中未提取到有效表格: {file_path}\n"
                f"  请确认文档包含表格或以 '姓名 总分 政治 ...' 格式书写的文本。"
            )

        result = pd.concat(frames, ignore_index=True)
        result = self._drop_unnamed_columns(result)
        print(f"  [DEBUG] Word 提取结果: {result.shape}")
        print(f"  [DEBUG] df.columns: {list(result.columns)}")
        return self._normalize(result)

    # ── 图片适配器 ────────────────────────────────────────────────────

    def load_image(self, file_path: str | Path) -> pd.DataFrame:
        """OCR 识别图片中的表格并返回 DataFrame.

        委托给 ``preprocess_image()``（PaddleOCR PPStructure 引擎）.

        Raises
        ------
        ImportError
            paddleocr 未安装.
        ValueError
            图片中未识别到表格.
        """
        df = preprocess_image(file_path)
        return self._normalize(df)

    # ── 标准化 ────────────────────────────────────────────────────────

    @staticmethod
    def _normalize(df: pd.DataFrame) -> pd.DataFrame:
        """将任意 DataFrame 标准化为统一 schema.

        先尝试精确匹配（``COLUMN_ALIASES``），
        再通过 ``normalize_columns()`` 做模糊匹配兜底。

        Returns
        -------
        pd.DataFrame
            ``["姓名", "总分", "政治", "英语", "业务课一", "业务课二", ...]``
        """
        if df.empty:
            return pd.DataFrame(columns=STANDARD_COLUMNS.copy())

        df = df.copy()
        # 全局清洗：删除 Unnamed 空列
        unnamed_drops = [
            c for c in df.columns
            if "unnamed" in str(c).lower() and df[c].isna().all()
        ]
        if unnamed_drops:
            df = df.drop(columns=unnamed_drops)

        # 表头清洗：去换行/空格/括号内容，独立数字1/2 → 一/二
        def _clean_col(c: object) -> str:
            s = str(c).strip().replace("\n", "").replace(" ", "").replace("\u3000", "")
            s = re.sub(r"[（(][^）)]*[）)]", "", s)  # 删除 (100分) 等后缀
            s = re.sub(r"(?<!\d)(\d)(?!\d)", lambda m: {"1": "一", "2": "二"}.get(m.group(1), m.group(1)), s)
            return s

        df.columns = [_clean_col(c) for c in df.columns]
        print(f"  [DEBUG] 清洗后列名: {list(df.columns)}")

        # 1. 精确匹配（现有逻辑）
        mapped: dict[str, str] = {}
        unmapped: list[str] = []

        for col in df.columns:
            lowered = col.lower()
            match = COLUMN_ALIASES.get(col) or COLUMN_ALIASES.get(lowered)
            if match:
                mapped[col] = match
            elif lowered in {v.lower() for v in COLUMN_ALIASES}:
                mapped[col] = col
            else:
                unmapped.append(col)

        # 2. 对未匹配列做模糊匹配兜底
        if unmapped:
            from difflib import get_close_matches

            ALL_ALIASES = sorted(set(_ALIAS_FLAT.keys()), key=len, reverse=True)
            newly_mapped: list[str] = []
            for col in list(unmapped):
                matches = get_close_matches(col, ALL_ALIASES, n=3, cutoff=0.80)
                if not matches:
                    matches = get_close_matches(col.lower(), ALL_ALIASES, n=3, cutoff=0.80)
                if matches:
                    best = _ALIAS_FLAT.get(matches[0]) or _ALIAS_FLAT.get(matches[0].lower())
                    # 禁止将非姓名列映射为姓名
                    if best == "姓名" and "姓名" not in col and "name" not in col.lower():
                        continue
                    if best:
                        mapped[col] = best
                        newly_mapped.append(col)
                        print(f"  [FUZZY] 列「{col}」→ {best}")
            for col in newly_mapped:
                unmapped.remove(col)

        # 3. 构建结果 DataFrame
        result = pd.DataFrame()

        for std_col in STANDARD_COLUMNS:
            src = next((src for src, dst in mapped.items() if dst == std_col), None)
            if src is not None:
                if std_col in ("总分", "政治", "英语", "业务课一", "业务课二"):
                    result[std_col] = pd.to_numeric(df[src], errors="coerce")
                else:
                    result[std_col] = df[src].astype(str)
            else:
                result[std_col] = pd.NA

        # 4. 追加未映射列（加 raw_ 前缀）
        for col in unmapped:
            result[f"raw_{col}"] = df[col]

        # 5. 核心科目缺失检查
        core = {"政治", "英语", "业务课一", "业务课二"}
        missing = core - set(result.columns)
        if missing:
            msg = f"未发现核心科目列: {missing}，请在侧边栏【数据结构修正】面板中手动指定列名。"
            result.attrs["_missing_core"] = msg

        return result

    # ── 内部辅助 ──────────────────────────────────────────────────────

    @staticmethod
    def _clean_cell(val: Any) -> str:
        """清洗 PDF 提取的表格单元格."""
        if val is None:
            return ""
        s = str(val).strip()
        return re.sub(r"\s+", " ", s)


def _looks_like_header(row: list[str]) -> bool:
    """判断 PDF 表格行是否为表头（而不是数据行）."""
    if not row:
        return False
    first = row[0].replace(" ", "").replace("\n", "")
    # 表头关键词
    header_kw = ["序号", "序", "名次", "排名", "编号", "学院", "专业", "科目", "姓名",
                 "政治", "外语", "英语", "业务课", "总分", "总成绩", "初试", "备注"]
    if any(kw in first for kw in header_kw):
        return True
    # 如果第一个单元格不是纯数字，且长度 > 1（排除序号），可能是表头
    if not first.isdigit() and len(first) > 1:
        return True
    return False


def _parse_free_text(text: str) -> pd.DataFrame | None:
    """从纯文本中提取 ``(姓名, 总分, 政治, 英语, 业务课一, 业务课二)`` 格式数据.

    匹配模式如：``张三 350 70 65 120 95``（空格或 Tab 分隔）
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    rows: list[list[str]] = []
    header_map = {v: k for k, v in COLUMN_ALIASES.items()}
    # 补充标准列名本身
    for c in STANDARD_COLUMNS:
        header_map[c] = c

    # 尝试按分隔符拆分
    for line in lines:
        parts = re.split(r"[\s\t|,，、]+", line)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 6:
            rows.append(parts[:6])
        elif len(parts) == 2 and parts[0] in STANDARD_COLUMNS:
            # 可能是一行一个字段，跳过
            continue

    if rows:
        return pd.DataFrame(rows, columns=STANDARD_COLUMNS)

    return None
